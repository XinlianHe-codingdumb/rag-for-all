from __future__ import annotations

import argparse
from pathlib import Path


PAGES = [
    (
        "Northstar Employee Handbook",
        "Learning allowance",
        "Every full-time teammate receives a SGD 1,200 learning allowance each calendar year. "
        "Manager approval is required before booking a course.",
    ),
    (
        "Time away",
        "Annual leave",
        "Employees receive 18 days of annual leave. Submit planned leave in the people portal "
        "at least two weeks in advance when possible.",
    ),
    (
        "Remote work",
        "Home office support",
        "The company reimburses up to SGD 500 for an ergonomic home-office setup. Keep the receipt "
        "and submit it within 30 days.",
    ),
]


def create_pdf(output: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

    output.parent.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    story = []
    for index, (title, heading, body) in enumerate(PAGES):
        story.extend(
            [
                Paragraph(title, styles["Title"]),
                Spacer(1, 20),
                Paragraph(heading, styles["Heading2"]),
                Spacer(1, 8),
                Paragraph(body, styles["BodyText"]),
            ]
        )
        if index < len(PAGES) - 1:
            story.append(PageBreak())
    SimpleDocTemplate(str(output), pagesize=A4, title="RAG FOR ALL parser fixture").build(story)


def create_docx(output: Path) -> None:
    from docx import Document
    from docx.enum.text import WD_BREAK

    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.core_properties.title = "RAG FOR ALL parser fixture"
    for index, (title, heading, body) in enumerate(PAGES):
        document.add_heading(title, 0)
        document.add_heading(heading, 1)
        document.add_paragraph(body)
        if index < len(PAGES) - 1:
            document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("format", choices=("pdf", "docx"))
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    if args.format == "pdf":
        create_pdf(args.output)
    else:
        create_docx(args.output)


if __name__ == "__main__":
    main()
